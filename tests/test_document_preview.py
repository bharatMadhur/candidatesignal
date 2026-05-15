from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from resume_intel.web import _docx_to_safe_html


class DocumentPreviewTests(unittest.TestCase):
    def test_docx_preview_escapes_user_content(self):
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            document = Document()
            document.add_paragraph("Senior Engineer <script>alert(1)</script>")
            document.save(path)

            rendered = _docx_to_safe_html(path)

        self.assertIn("Senior Engineer", rendered)
        self.assertIn("&lt;script&gt;", rendered)
        self.assertNotIn("<script>", rendered)

    def test_docx_preview_preserves_basic_resume_structure(self):
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            document = Document()
            document.add_heading("Candidate Name", level=1)
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            run = paragraph.add_run("Led AI platform")
            run.bold = True
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Skill"
            table.rows[0].cells[1].text = "Years"
            table.rows[1].cells[0].text = "Python"
            table.rows[1].cells[1].text = "6"
            document.save(path)

            rendered = _docx_to_safe_html(path)

        self.assertIn("<h2>Candidate Name</h2>", rendered)
        self.assertIn('<li class="bullet"><strong>Led AI platform</strong></li>', rendered)
        self.assertIn("<th>Skill</th>", rendered)
        self.assertIn("<td>Python</td>", rendered)


if __name__ == "__main__":
    unittest.main()
