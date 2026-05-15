from __future__ import annotations

import hashlib
import re
import shlex
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageOps
import pypdfium2 as pdfium
from pypdf import PdfReader

from .settings import Settings


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".tif", ".tiff", ".bmp"}
BARE_DOMAIN_RE = re.compile(
    r"^(?:[a-z0-9](?:[a-z0-9-]{0,61}[a-z0-9])?\.)+"
    r"(?:com|ai|io|dev|co|net|org|me|app|xyz|in|us|edu)"
    r"(?:/[^\s<>()]*)?$",
    re.I,
)


@dataclass(frozen=True)
class ExtractedDocument:
    document_id: str
    source_file: Path
    text: str
    method: str
    page_count: int | None = None
    pages: list[dict] | None = None
    links: list[dict] | None = None


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()[:16]


def extract_document(path: Path, settings: Settings, work_dir: Path) -> ExtractedDocument:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path, settings, work_dir)
    if suffix in IMAGE_EXTENSIONS:
        return _extract_image(path, settings, work_dir)
    if suffix in {".txt", ".md"}:
        text = path.read_text(errors="ignore")
        links = _extract_inline_links(text, "text", 1)
        return ExtractedDocument(file_hash(path), path, text, "text", 1, [{"page_number": 1, "text": text, "method": "text", "quality_flags": []}], links)
    raise ValueError(f"Unsupported file type: {path}")


def _extract_docx(path: Path) -> ExtractedDocument:
    doc = Document(path)
    chunks: list[str] = []
    links: list[dict] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
        links.extend(_docx_paragraph_hyperlinks(paragraph))
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    links.extend(_docx_paragraph_hyperlinks(paragraph))
    links = _dedupe_links([*links, *_docx_relationship_hyperlinks(doc)])
    text = _append_extracted_links("\n".join(chunks), links)
    quality_flags = ["docx_hyperlinks_found"] if links else []
    return ExtractedDocument(file_hash(path), path, text, "docx", 1, [{"page_number": 1, "text": text, "method": "docx", "quality_flags": quality_flags}], links)


def _extract_pdf(path: Path, settings: Settings, work_dir: Path) -> ExtractedDocument:
    reader = PdfReader(str(path))
    links = _extract_pdf_annotation_links(reader)
    page_texts: list[str] = []
    pages: list[dict] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_links = [link for link in links if link.get("page_number") == page_number]
        page_text = _append_extracted_links(text.strip(), page_links)
        if page_text:
            page_texts.append(f"[PAGE {page_number}]\n{page_text}")
        quality_flags = _quality_flags(text, "pdf_text", len(reader.pages))
        if page_links:
            quality_flags = sorted(set([*quality_flags, "pdf_annotation_links_found"]))
        pages.append(
            {
                "page_number": page_number,
                "text": page_text,
                "method": "pdf_text" if text.strip() else "pdf_empty",
                "quality_flags": quality_flags,
                "links": page_links,
            }
        )

    text = "\n\n".join(page_texts).strip()
    if _looks_like_text_pdf(text, len(reader.pages)):
        return ExtractedDocument(file_hash(path), path, text, "pdf_text", len(reader.pages), pages, links)

    if settings.ocr_mode in {"external", "remote"}:
        return _extract_scanned_pdf(path, settings, work_dir, len(reader.pages), links)

    for page in pages:
        page["quality_flags"] = sorted(set([*page["quality_flags"], "scanned_or_low_text"]))
    return ExtractedDocument(file_hash(path), path, text, "pdf_empty_or_scanned", len(reader.pages), pages, links)


def _extract_image(path: Path, settings: Settings, work_dir: Path) -> ExtractedDocument:
    if settings.ocr_mode not in {"external", "remote"}:
        raise ValueError("Image resumes require OCR_MODE=external or OCR_MODE=remote")

    normalized_path = _normalize_image_for_ocr(path, work_dir)
    method = "image_remote_ocr" if settings.ocr_mode == "remote" else "image_external_ocr"
    ocr_text = _run_ocr_pages(settings, [normalized_path]).strip()
    links = _extract_inline_links(ocr_text, "image_ocr", 1)
    text = _append_extracted_links(ocr_text, links)
    quality_flags = _quality_flags(ocr_text, method, 1)
    quality_flags = sorted(set([*quality_flags, "normalized_image_for_ocr"]))
    return ExtractedDocument(
        file_hash(path),
        path,
        text,
        method,
        1,
        [{"page_number": 1, "text": text, "method": method, "quality_flags": quality_flags, "links": links}],
        links,
    )


def _normalize_image_for_ocr(path: Path, work_dir: Path) -> Path:
    image_dir = work_dir / file_hash(path) / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    normalized_path = image_dir / "page_001.png"
    with Image.open(path) as image:
        normalized = ImageOps.exif_transpose(image)
        if normalized.mode not in {"RGB", "L"}:
            normalized = normalized.convert("RGB")
        normalized.save(normalized_path, "PNG")
    return normalized_path


def _looks_like_text_pdf(text: str, page_count: int) -> bool:
    min_chars = max(200, page_count * 120)
    return len(text) >= min_chars


def _extract_scanned_pdf(
    path: Path, settings: Settings, work_dir: Path, page_count: int, links: list[dict] | None = None
) -> ExtractedDocument:
    image_dir = work_dir / file_hash(path) / "pages"
    image_dir.mkdir(parents=True, exist_ok=True)
    pdf = pdfium.PdfDocument(str(path))
    scale = settings.pdf_render_dpi / 72
    image_paths: list[str] = []
    for index, page in enumerate(pdf, start=1):
        image = page.render(scale=scale).to_pil()
        image_path = image_dir / f"page_{index:03d}.png"
        image.save(image_path)
        image_paths.append(str(image_path))

    method = "pdf_remote_ocr" if settings.ocr_mode == "remote" else "pdf_external_ocr"
    ocr_output = _run_ocr_pages(settings, [Path(item) for item in image_paths])
    links = links or []
    text = _append_extracted_links(ocr_output.strip(), links)
    page_chunks = _split_ocr_pages(ocr_output.strip())
    pages: list[dict] = []
    for index, chunk in enumerate(page_chunks):
        page_number = index + 1
        page_links = [link for link in links if link.get("page_number") == page_number]
        quality_flags = _quality_flags(chunk, method, page_count)
        if page_links:
            quality_flags = sorted(set([*quality_flags, "pdf_annotation_links_found"]))
        pages.append(
            {
                "page_number": page_number,
                "text": _append_extracted_links(chunk, page_links),
                "method": method,
                "quality_flags": quality_flags,
                "links": page_links,
            }
        )
    return ExtractedDocument(
        file_hash(path),
        path,
        text,
        method,
        page_count,
        pages,
        links,
    )


def _run_ocr_pages(settings: Settings, image_paths: list[Path]) -> str:
    if settings.ocr_mode == "remote":
        return _run_remote_ocr_pages(settings, image_paths)
    if not settings.ocr_command:
        raise ValueError("OCR_MODE=external requires OCR_COMMAND to be set")
    command = [*shlex.split(settings.ocr_command), *[str(path) for path in image_paths]]
    completed = subprocess.run(
        command,
        check=True,
        text=True,
        capture_output=True,
    )
    return completed.stdout


def _run_remote_ocr_pages(settings: Settings, image_paths: list[Path]) -> str:
    if not settings.ocr_remote_url:
        raise ValueError("OCR_MODE=remote requires OCR_REMOTE_URL to be set")
    endpoint = settings.ocr_remote_url.rstrip("/")
    if not endpoint.endswith("/ocr/pages"):
        endpoint = f"{endpoint}/ocr/pages"
    headers = {}
    if settings.ocr_remote_auth == "google_id_token":
        headers["Authorization"] = f"Bearer {_google_identity_token(settings.ocr_remote_audience or settings.ocr_remote_url)}"
        if settings.ocr_remote_token:
            headers["X-OCR-Token"] = settings.ocr_remote_token
    elif settings.ocr_remote_token:
        headers["Authorization"] = f"Bearer {settings.ocr_remote_token}"
    files = []
    handles = []
    try:
        for path in image_paths:
            handle = path.open("rb")
            handles.append(handle)
            files.append(("files", (path.name, handle, "image/png")))
        response = httpx.post(endpoint, headers=headers, files=files, timeout=settings.ocr_remote_timeout_seconds)
        response.raise_for_status()
    finally:
        for handle in handles:
            handle.close()
    payload = response.json()
    pages = payload.get("pages") or []
    chunks: list[str] = []
    for index, page in enumerate(pages, start=1):
        text = str(page.get("text") or "").strip()
        if text:
            page_number = int(page.get("page_number") or index)
            chunks.append(f"[PAGE {page_number}]\n{text}")
    return "\n\n".join(chunks)


def _google_identity_token(audience: str | None) -> str:
    if not audience:
        raise ValueError("OCR_REMOTE_AUTH=google_id_token requires OCR_REMOTE_AUDIENCE or OCR_REMOTE_URL")
    try:
        import google.auth.transport.requests
        import google.oauth2.id_token
    except ImportError as exc:  # pragma: no cover - depends on GCP deployment dependency
        raise RuntimeError("Google identity-token OCR auth requires google-auth") from exc
    request = google.auth.transport.requests.Request()
    return google.oauth2.id_token.fetch_id_token(request, audience.rstrip("/"))


def _split_ocr_pages(text: str) -> list[str]:
    chunks = [chunk.strip() for chunk in text.split("\f") if chunk.strip()]
    if len(chunks) <= 1 and "[PAGE " in text:
        chunks = [chunk.strip() for chunk in re.split(r"(?=\[PAGE\s+\d+\])", text) if chunk.strip()]
    return chunks or ([text] if text else [])


def _quality_flags(text: str, method: str, page_count: int | None = None) -> list[str]:
    flags: list[str] = []
    if method.startswith("pdf") and len((text or "").strip()) < 120:
        flags.append("low_text_density")
    if method in {"pdf_external_ocr", "pdf_remote_ocr"} and not (text or "").strip():
        flags.append("ocr_empty")
    if method in {"image_external_ocr", "image_remote_ocr"} and not (text or "").strip():
        flags.append("ocr_empty")
    if method in {"image_external_ocr", "image_remote_ocr"}:
        flags.append("image_ocr")
    if method == "pdf_text" and page_count and len((text or "").strip()) < max(120, page_count * 80):
        flags.append("possible_layout_loss")
    return flags


def _extract_inline_links(text: str, source: str, page_number: int | None = None) -> list[dict]:
    links: list[dict] = []
    for token in text.replace("(", " ").replace(")", " ").replace("<", " ").replace(">", " ").split():
        normalized = _normalize_extracted_url(token)
        if normalized:
            links.append(
                {
                    "url": normalized,
                    "label": _label_for_url(normalized),
                    "page_number": page_number,
                    "source": source,
                }
            )
    return _dedupe_links(links)


def _extract_pdf_annotation_links(reader: PdfReader) -> list[dict]:
    links: list[dict] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            annotations = page.get("/Annots") or []
        except Exception:
            annotations = []
        for annotation_ref in annotations:
            annotation = _resolve_pdf_object(annotation_ref)
            link = _pdf_annotation_link(annotation, page_number)
            if link:
                links.append(link)
    return _dedupe_links(links)


def _pdf_annotation_link(annotation: Any, page_number: int) -> dict | None:
    if not hasattr(annotation, "get"):
        return None
    action = _resolve_pdf_object(annotation.get("/A"))
    uri = action.get("/URI") if hasattr(action, "get") else None
    if not uri:
        uri = annotation.get("/URI")
    url = _normalize_extracted_url(str(uri)) if uri else ""
    if not url:
        return None
    return {
        "url": url,
        "label": _pdf_annotation_label(annotation, url),
        "page_number": page_number,
        "source": "pdf_annotation",
    }


def _pdf_annotation_label(annotation: Any, url: str) -> str:
    for key in ("/Contents", "/T", "/TU", "/Alt"):
        try:
            value = annotation.get(key)
        except Exception:
            value = None
        if value:
            label = str(value).strip()
            if label:
                return label[:120]
    return _label_for_url(url)


def _docx_paragraph_hyperlinks(paragraph: Any) -> list[dict]:
    links: list[dict] = []
    try:
        hyperlink_nodes = paragraph._p.xpath(".//w:hyperlink")
    except Exception:
        return links
    for hyperlink in hyperlink_nodes:
        rel_id = hyperlink.get(qn("r:id"))
        if not rel_id or rel_id not in paragraph.part.rels:
            continue
        target = _normalize_extracted_url(paragraph.part.rels[rel_id].target_ref)
        if not target:
            continue
        label = _docx_hyperlink_label(hyperlink) or _label_for_url(target)
        links.append({"url": target, "label": label, "page_number": 1, "source": "docx_hyperlink"})
    return links


def _docx_relationship_hyperlinks(doc: Document) -> list[dict]:
    links: list[dict] = []
    for relationship in doc.part.rels.values():
        reltype = str(getattr(relationship, "reltype", ""))
        if "hyperlink" not in reltype:
            continue
        url = _normalize_extracted_url(getattr(relationship, "target_ref", ""))
        if url:
            links.append({"url": url, "label": _label_for_url(url), "page_number": 1, "source": "docx_relationship"})
    return links


def _docx_hyperlink_label(hyperlink: Any) -> str:
    try:
        return " ".join(text.text or "" for text in hyperlink.xpath(".//w:t")).strip()
    except Exception:
        return ""


def _resolve_pdf_object(value: Any) -> Any:
    if hasattr(value, "get_object"):
        try:
            return value.get_object()
        except Exception:
            return value
    return value


def _append_extracted_links(text: str, links: list[dict] | None) -> str:
    link_lines = _extracted_link_lines(links or [])
    if not link_lines:
        return text.strip()
    body = text.strip()
    link_text = "\n".join(["[EXTRACTED DOCUMENT LINKS]", *link_lines])
    return f"{body}\n\n{link_text}".strip() if body else link_text


def _extracted_link_lines(links: list[dict]) -> list[str]:
    lines: list[str] = []
    for link in _dedupe_links(links):
        url = str(link.get("url") or "").strip()
        if not url:
            continue
        label = str(link.get("label") or _label_for_url(url)).strip()
        page_number = link.get("page_number")
        source = str(link.get("source") or "document_link").upper()
        prefix = f"[PAGE {page_number} {source}]" if page_number else f"[{source}]"
        lines.append(f"{prefix} {label}: {url}")
    return lines


def _normalize_extracted_url(value: str) -> str:
    token = (value or "").strip().strip(".,;:)]}>\"'")
    if not token:
        return ""
    lower = token.lower()
    if lower.startswith("mailto:"):
        return token
    if lower.startswith(("http://", "https://")):
        return token
    if lower.startswith("www."):
        return f"https://{token}"
    if any(domain in lower for domain in ("linkedin.com/", "github.com/", "behance.net/", "dribbble.com/", "medium.com/")):
        return f"https://{token}" if "://" not in token else token
    if "@" not in token and BARE_DOMAIN_RE.match(token):
        return f"https://{token}"
    return ""


def _label_for_url(url: str) -> str:
    lower = url.lower()
    if "linkedin.com/" in lower:
        return "LinkedIn"
    if "github.com/" in lower:
        return "GitHub"
    if "behance.net/" in lower:
        return "Behance"
    if "dribbble.com/" in lower:
        return "Dribbble"
    if lower.startswith("mailto:"):
        return "Email"
    return "Portfolio"


def _dedupe_links(links: list[dict]) -> list[dict]:
    deduped: list[dict] = []
    seen: set[tuple[str, str, int | None]] = set()
    for link in links:
        url = str(link.get("url") or "").strip()
        if not url:
            continue
        label = str(link.get("label") or _label_for_url(url)).strip()
        page_number = link.get("page_number") if isinstance(link.get("page_number"), int) else None
        source = str(link.get("source") or "document_link")
        key = (url.lower(), label.lower(), page_number)
        if key in seen:
            continue
        seen.add(key)
        deduped.append({"url": url, "label": label, "page_number": page_number, "source": source})
    return deduped
